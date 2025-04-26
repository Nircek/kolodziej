#!/usr/bin/env python3
# -*- coding: UTF-8 -*-

# inspirated by http://people.cas.uab.edu/~mosya/cl/CPPcircle.html
# copyright (c) 2011-2014 Nikolai Chernov
# edited by Nircek 2019

# MIT License

# Copyright (c) 2019-2020, 2025 Nircek

# Permission is hereby granted, free of charge, to any person obtaining a copy
# of this software and associated documentation files (the "Software"), to deal
# in the Software without restriction, including without limitation the rights
# to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
# copies of the Software, and to permit persons to whom the Software is
# furnished to do so, subject to the following conditions:

# The above copyright notice and this permission notice shall be included in all
# copies or substantial portions of the Software.

# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
# IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
# FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
# AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
# LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
# OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
# SOFTWARE.

from math import sin, cos, pi, atan2, sqrt
import os
import io
import sys
import traceback
from tkinter.simpledialog import askstring
from tkinter import Tk, Menu, Text, INSERT, mainloop
from tkinter import filedialog, messagebox  # Tk must be installed
from webbrowser import open_new
from docx.shared import Pt  # , RGBColor

# https://stackoverflow.com/a/57189298/6732111 pylint: disable=E0611,E1101
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Cm
import docx  # pip install python-docx
from PIL import Image, ImageDraw, ImageFont  # ImageTk

print("Starting... ", end="", flush=True)
# pylint: disable=C0301,C0114,C0103,C0115,C0116,R0913,R0914,R0903,R0912,R0915,W0703

TEAM = "J. Kowalski"
VERSION = "v2.0.0"
DATE = "26.04.2025"
YEARS = "2019-2020, 2025"


class Data:
    def __init__(self, n=0, xs=None, ys=None):
        self.n = n
        if xs is not None and ys is not None:
            assert len(xs) == len(ys) == n, "not as much xs or ys as declared"
            self.X = xs
            self.Y = ys
        else:
            for _ in self.n:
                self.X += [0.0]
                self.Y += [0.0]

    @property
    def mX(self):
        return sum(self.X) / len(self.X)

    @property
    def mY(self):
        return sum(self.Y) / len(self.Y)

    def center(self):
        self.X = list(map(lambda x: x - self.mX, self.X))
        self.Y = list(map(lambda x: x - self.mY, self.Y))

    def scale(self):
        sXX = 0.0
        sYY = 0.0
        for i in range(self.n):
            sXX += self.X[i] * self.X[i]
            sYY += self.Y[i] * self.Y[i]
        scaling = sqrt((sXX + sYY) / self.n / 2)
        for _ in range(self.n):
            self.X[i] /= scaling
            self.Y[i] /= scaling

    def __repr__(self):
        s = ""
        s += f"DATA(len:{self.n})["
        for i in range(len(self.n)):
            s += f"({self.X[i]},{self.Y[i]}), "
        if self.n == 0:
            s = s[:-2]
        s += "]"
        return s


class Circle:
    def __init__(self, aa=0.0, bb=0.0, rr=1.0):
        self.a = aa
        self.b = bb
        self.r = rr
        self.s = 0.0
        self.i = 0
        self.j = 0

    def __repr__(self):
        s = "CIRCLE("
        s += "x:" + str(self.a)
        s += ", y:" + str(self.b)
        s += ", r:" + str(self.r)
        s += ", s:" + str(self.s)
        s += ", g:" + str(self.g)
        s += ", i:" + str(self.i)
        s += ", j:" + str(self.j)
        s += ")"
        return s


def SQR(t):
    return t * t


def Sigma(data, circle):
    summ = 0.0
    for i in range(data.n):
        dx = data.X[i] - circle.a
        dy = data.Y[i] - circle.b
        summ += SQR(sqrt(dx * dx + dy * dy) - circle.r)
    return sqrt(summ / data.n)


# , circle):
def CircleFitByLevenbergMarquardtFull(data, circleIni, LambdaIni):
    """
    Algorithm:  Levenberg-Marquardt running over the full parameter space (a,b,r)

    See a detailed description in Section 4.5 of the book by Nikolai Chernov:
    "Circular and linear regression: Fitting circles and lines by least squares"
    Chapman & Hall/CRC, Monographs on Statistics and Applied Probability, volume 117, 2010.

    Nikolai Chernov,  February 2014
    see the unedited code: http://people.cas.uab.edu/~mosya/cl/CircleFitByLevenbergMarquardtFull.cpp
    port from C++ to Python by Nircek in January 2019
    """
    IterMAX = 99
    factorUp = 10.0
    factorDown = 0.04
    ParLimit = 1.0e6
    epsilon = 3.0e-8
    New = circleIni
    New.s = Sigma(data, New)
    lambd = LambdaIni
    iterr = inner = 0
    code = -1
    while 1:
        Old = New
        iterr += 1
        if iterr > IterMAX:
            code = 1
            break
        Mu = Mv = Muu = Mvv = Muv = Mr = 0.0
        for i in range(data.n):
            dx = data.X[i] - Old.a
            dy = data.Y[i] - Old.b
            ri = sqrt(dx * dx + dy * dy)
            u = dx / ri
            v = dy / ri
            Mu += u
            Mv += v
            Muu += u * u
            Mvv += v * v
            Muv += u * v
            Mr += ri
        Mu /= data.n
        Mv /= data.n
        Muu /= data.n
        Mvv /= data.n
        Muv /= data.n
        Mr /= data.n
        F1 = Old.a + Old.r * Mu - data.mX
        F2 = Old.b + Old.r * Mv - data.mY
        F3 = Old.r - Mr
        Old.g = New.g = sqrt(F1 * F1 + F2 * F2 + F3 * F3)
        while 1:
            UUl = Muu + lambd
            VVl = Mvv + lambd
            Nl = 1.0 + lambd
            G11 = sqrt(UUl)
            G12 = Muv / G11
            G13 = Mu / G11
            G22 = sqrt(VVl - G12 * G12)
            G23 = (Mv - G12 * G13) / G22
            G33 = sqrt(Nl - G13 * G13 - G23 * G23)
            D1 = F1 / G11
            D2 = (F2 - G12 * D1) / G22
            D3 = (F3 - G13 * D1 - G23 * D2) / G33
            dR = D3 / G33
            dY = (D2 - G23 * dR) / G22
            dX = (D1 - G12 * dY - G13 * dR) / G11
            if (abs(dR) + abs(dX) + abs(dY)) / (1.0 + Old.r) < epsilon:
                code = 0
                break
            New.a = Old.a - dX
            New.b = Old.b - dY
            if abs(New.a) > ParLimit or abs(New.b) > ParLimit:
                code = 3
                break
            New.r = Old.r - dR
            if New.r <= 0.0:
                lambd *= factorUp
                inner += 1
                if inner > IterMAX:
                    code = 2
                break
            New.s = Sigma(data, New)
            if New.s < Old.s:
                lambd *= factorDown
            else:
                inner += 1
                if inner > IterMAX:
                    code = 2
                else:
                    lambd *= factorUp
            break
        if code != -1:
            break
    Old.i = iterr
    Old.j = inner
    return code, Old


# src: https://stackoverflow.com/a/39885430/6732111
def resource_path(relative_path):
    """Get absolute path to resource, works for dev and for PyInstaller"""
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS  # pylint: disable=W0212
    except Exception:
        base_path = os.environ.get("_MEIPASS2", os.path.abspath("."))
    return os.path.join(base_path, relative_path)


def loadFile(log, fn):
    f = open(fn, "r")
    Xs = []
    Ys = []
    for ff in f:
        ffs = ff.replace(",", ".").split(maxsplit=2)
        if len(ffs) < 2:
            log += f"WARN: ignoring '{ff}'\n"
            continue
        if len(ffs) > 2:
            log += f"WARN: ignoring '{ffs[2]}'\n"
        x = float(ffs[0])
        y = float(ffs[1])
        Xs += [x]
        Ys += [y]
    return log, Xs, Ys


def calcCircle(Xs, Ys):
    LambdaIni = 0.001
    data1 = Data(len(Xs), Xs, Ys)
    circleIni = Circle()
    code, circle = CircleFitByLevenbergMarquardtFull(data1, circleIni, LambdaIni)
    return code, circle.a, circle.b, circle.r, circle.s, circle.i, data1


def makeCircle(cx, cy, cr, W, M, x, y, r, imgd, fnt, fntb, circ, fit_circle, north_sign):
    if fit_circle:
        circ(cx, cy, 8, {"fill": "red"})
        circ(cx, cy, cr)
    imgd.line((0, y(0) + M, W, y(0) + M), fill="black")
    imgd.line((W - 15, y(0) - 5 + M, W, y(0) + M, W - 15, y(0) + 5 + M), fill="black")
    imgd.line((x(0), M, x(0), W + M), fill="black")
    imgd.line((x(0) - 5, 15 + M, x(0), M, x(0) + 5, 15 + M), fill="black")
    imgd.text((x(0), 0), "Y", font=fntb, fill="black")
    imgd.text((W, y(0)), "X", font=fntb, fill="black")
    scale = r(1000)
    for ii in range(1, 10):
        l = 12
        w = 0
        if ii == 5:
            l = 20
            w = 2
        imgd.line([16 + ii * scale / 10, W - l + M, 16 + ii * scale / 10, W - 2 + M], fill="black", width=w)
    imgd.line([16, W - 18 + M, 16, W - 2 + M, 16 + scale, W - 2 + M, 16 + scale, W - 18 + M], fill="black")
    imgd.text((16 - 6, W - 18 - 42 + M), "0", font=fnt, fill="black")
    imgd.text((16 + scale - 50, W - 18 - 42 + M), "1000 mm", font=fnt, fill="black")
    if north_sign:
        imgd.line(
            [W + M - 64, W + M - 55, W + M - 64, W + M - 205, W + M - 34, W + M - 55, W + M - 34, W + M - 205],
            fill="black",
            width=3,
        )
        imgd.line(
            [W + M - 49, W + M, W + M - 49, W + M - 410, W + M - 64, W + M - 250, W + M - 34, W + M - 250],
            fill="black",
            width=3,
        )


def isBlank(imgpx, m, sf, ix, iy):
    for dx in range(int(m * sf)):
        for dy in range(sf):
            try:
                if imgpx[ix + dx, iy + dy] != (255, 255, 255, 0):
                    return False
            except IndexError:
                return False
    return True


def findPlace(imgpx, e, angle, ex, ey):
    for r in (22, 33, 44, 55, 66, 77, 88, 99):
        sf = 54 if e[3] else 30  # size of font + 6
        for a in range(24):
            a = a / 12 * pi + angle
            m = len(e[2]) / 2
            ix = int(ex + r * sin(a) - m * sf / 2)
            iy = int(ey + r * cos(a) - sf / 2)
            if isBlank(imgpx, m, sf, ix, iy):
                return ix, iy
    return None


def makePoints(imgd, imgpx, cx, cy, cr, arr, fnt, fntb):  # , circ):
    for e in arr:
        # r = 33  # radius from point
        ex, ey = e[0], e[1]
        angle = atan2(ex - cx, ey - cr)
        if cr > sqrt((ex - cx) ** 2 + (ey - cy) ** 2):
            angle += pi
        pl = findPlace(imgpx, e, angle, ex, ey)
        if pl is None:
            print('ERR: no place for "', e[2], '"', sep="")
        else:
            imgd.text((pl[0], pl[1]), e[2], font=(fntb if e[3] else fnt), fill=("red" if e[3] else "black"))


def genImage(X, Y, Wc, data1, ca, cb, crr, fit_circle, north_sign, show_numbers):
    W = 2048  # in px
    M = 36
    x, y = lambda x: (x - X) * W / Wc, lambda y: -(y - Y - Wc) * W / Wc

    def r(x):
        return x * W / Wc

    img = Image.new("RGBA", (W + M, W + M), (255, 255, 255, 0))
    imgd = ImageDraw.Draw(img)
    fnt = ImageFont.truetype(resource_path("./fonts/Roboto-Regular.ttf"), 24)
    fntb = ImageFont.truetype(resource_path("./fonts/Roboto-Regular.ttf"), 48)
    imgpx = img.load()
    arr = [(x(data1.X[i]), y(data1.Y[i]) + M, str(i + 1), False) for i in range(data1.n)]
    cx, cy, cr = x(ca), y(cb) + M, r(crr)
    circ = (
        lambda x, y, r, a={"outline": "red"}, i=None: imgd.ellipse((x - r, y - r, x + r, y + r), **a)
        if i is None
        else imgd.arc((x - r, y - r, x + r, y + r), i[0], i[1], **a)
    )
    makeCircle(cx, cy, cr, W, M, x, y, r, imgd, fnt, fntb, circ, fit_circle, north_sign)
    for e in arr:
        circ(e[0], e[1], 5, {"outline": "black"})
        circ(e[0], e[1], 4, {"outline": "black"})
        circ(e[0], e[1], 3, {"outline": "black"})
    if fit_circle:
        arr += [(cx, cy, "S", True)]
    if show_numbers:
        makePoints(imgd, imgpx, cx, cy, cr, arr, fnt, fntb)
    return img


def str_comma(i):
    return str(i).replace(".", ",")


def b(s, l):
    return str(s)[:l] + (l - len(str(s))) * " "


def handleFile(X, Y, W, log, doc, i, fn, t, n, fit_circle, north_sign, show_numbers):
    print(f"[{i + 1}/{n}]FILE: {fn}\nCalculating... ", end="", flush=True)
    if i:
        doc.add_page_break()
    log, Xs, Ys = loadFile(log, fn)
    code, ca, cb, cr, cs, ci, data1 = calcCircle(Xs, Ys)
    print("DONE")
    log += b(i + 1, t[0]) + b(fn, t[1])
    if code in (1, 2):
        log += "ERR: Iterations maxed out.\n"
    elif code == 3:
        log += "ERR: Fitting circle too big.\n"
    elif code == 0:
        print("Making a chart... ", end="", flush=True)
        log += (
            b(str_comma(ca), t[2])
            + b(str_comma(cb), t[3])
            + b(str_comma(cr), t[4])
            + b(str_comma(cs), t[5])
            + str(str_comma(ci))
            + "\n"
        )
        img = genImage(X, Y, W, data1, ca, cb, cr, fit_circle, north_sign, show_numbers)
        r = doc.add_paragraph().add_run("Przekrój: ")
        r.bold = True
        r.font.size = Pt(16)
        doc.add_paragraph().add_run("Głębokość:  m").font.size = Pt(14)
        with io.BytesIO() as out:
            img.save(out, format="PNG")
            p = doc.add_paragraph()
            p.add_run().add_picture(out, width=Cm(17))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if fit_circle:
            p = doc.add_paragraph("\tWspółrzędne środka okręgu:\n\tX")
            p.paragraph_format.tab_stops.add_tab_stop(Cm(9), WD_TAB_ALIGNMENT.LEFT)
            p.add_run("S").font.subscript = True
            p.add_run(f" = {round(ca)} mm Y")
            p.add_run("S").font.subscript = True
            p.add_run(f" = {round(cb)} mm")
            doc.add_paragraph(f"\tŚrednica okręgu:\n\tD = {round(cr * 2)} mm").paragraph_format.tab_stops.add_tab_stop(
                Cm(9), WD_TAB_ALIGNMENT.LEFT
            )
        else:
            doc.add_paragraph("\n")
            doc.add_paragraph("\n")
        doc.add_paragraph().add_run(f"Data pomiaru: {DATE} r.\nZespół pomiarowy: {TEAM}").font.size = Pt(7)
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.right_margin = Cm(0.25)
        print("DONE")
    else:
        log += f"Unexpected code: {code}\n"
    return log, doc


def makeWindow(doc, log, width):
    tk = Tk()
    tk.title(f"Kołodziej {VERSION}")
    menubar = Menu(tk)
    filemenu = Menu(menubar, tearoff=0)

    def saveas():
        return (lambda x: doc.save(x) if x else "")(
            (
                lambda: filedialog.asksaveasfilename(
                    defaultextension=".docx", filetypes=(("docx files", "*.docx"), ("all files", "*.*"))
                )
            )()
        )

    filemenu.add_command(label="Save as...", command=saveas)
    filemenu.add_command(label="Exit", command=tk.destroy)
    menubar.add_cascade(label="File", menu=filemenu)
    helpmenu = Menu(menubar, tearoff=0)
    helpmenu.add_command(label="Get source code", command=lambda: open_new("https://github.com/Nircek/kolodziej"))
    helpmenu.add_command(
        label="About...",
        command=lambda: messagebox.showinfo(
            f"Kołodziej {VERSION}",
            f"Kołodziej {VERSION} by Nircek\nCopyright \N{COPYRIGHT SIGN} Nircek {YEARS}\nLAST EDIT {DATE}",
        ),
    )
    menubar.add_cascade(label="Help", menu=helpmenu)
    tk.config(menu=menubar)
    tk.bind("<Control-s>", lambda x: saveas())
    w = Text(tk, width=width, foreground="black")
    w.insert(INSERT, log)
    w.pack()
    w.configure(state="disabled")
    w.configure(background="white")
    w.configure(inactiveselectbackground=w.cget("selectbackground"))
    mainloop()


def askintegerdef(title, msg, defa):
    r = askstring(title, msg)
    try:
        r = int(r)
    except ValueError:
        r = defa
    return r


def main():
    try:
        log = ""
        files = sys.argv[1:]
        tk = Tk()
        tk.withdraw()
        print("DONE")
        if not files:
            print("Choosing files... ", end="", flush=True)
            files = filedialog.askopenfilenames(filetypes=(("txt files", "*.txt"), ("all files", "*.*")))
            if not files:
                files = []
            print("DONE")
        files = list(files)
        msg = "Please enter the {} coordinate of the lower-left corner (default: -5000):"
        X = askintegerdef("Input Required", msg.format("X"), -5000)
        Y = askintegerdef("Input Required", msg.format("Y"), -5000)
        msg = "Enter the width of the image (recommended: slightly larger than the circle diameter, default: 10000):"
        W = askintegerdef("Specify Width", msg, 10000)
        msg = "Would you like to fit the points into a circle? (default: Yes)"
        fit_circle = messagebox.askyesno("Circle Fitting", msg)
        north_sign = messagebox.askyesno("North Indicator", "Would you like to draw a north indicator? (default: Yes)")
        msg = "Would you like to display point numbers on the image? (default: Yes)"
        show_numbers = messagebox.askyesno("Point Labels", msg)
        tk.destroy()
        t = [
            len(str(len(files))) + 1,
            max([len(x) for x in files + ["Name"]]) + 1,
            len(str(-0.0001 / 3)) + 1,
            len(str(-0.0001 / 3)) + 1,
            len(str(-0.0001 / 3)) + 1,
            len(str(-0.0001 / 3)) + 1,
            len("Iterations"),
        ]
        log += b("I", t[0]) + b("Name", t[1]) + b("X", t[2]) + b("Y", t[3])
        log += b("Radius", t[4]) + b("Sigma", t[5]) + "Iterations" + "\n"
        doc = docx.Document()
        for i, fn in enumerate(files):
            log, doc = handleFile(X, Y, W, log, doc, i, fn, t, len(files), fit_circle, north_sign, show_numbers)
        makeWindow(doc, log, sum(t))
    except Exception as _:
        print(traceback.format_exc(), file=sys.stderr)
        messagebox.showerror("Fatal error", traceback.format_exc())


if __name__ == "__main__":
    main()
